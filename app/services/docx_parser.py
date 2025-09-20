"""
DOCX parsing service using python-docx and docx2txt (approved packages)
"""
from docx import Document  # python-docx
import docx2txt
from pathlib import Path
from typing import Dict, Optional
import logging


class DOCXParser:
    """Parse DOCX files to extract text content"""

    def __init__(self):
        self.logger = logging.getLogger(__name__)

    def extract_text_python_docx(self, file_path: str) -> Dict[str, any]:
        """Extract text using python-docx (primary method)"""
        try:
            doc = Document(file_path)
            text_content = ""
            metadata = {
                'total_paragraphs': 0,
                'total_tables': 0,
                'method_used': 'python-docx',
                'success': False,
                'error': None
            }

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n"
                    metadata['total_paragraphs'] += 1

            # Extract text from tables
            for table in doc.tables:
                metadata['total_tables'] += 1
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content += cell.text + " "
                    text_content += "\n"

            if text_content.strip():
                metadata['success'] = True
                return {
                    'text': text_content.strip(),
                    'metadata': metadata
                }
            else:
                metadata['error'] = "No text extracted"
                return {'text': "", 'metadata': metadata}

        except Exception as e:
            self.logger.error(f"python-docx extraction failed: {str(e)}")
            return {
                'text': "",
                'metadata': {
                    'total_paragraphs': 0,
                    'total_tables': 0,
                    'method_used': 'python-docx',
                    'success': False,
                    'error': str(e)
                }
            }

    def extract_text_docx2txt(self, file_path: str) -> Dict[str, any]:
        """Extract text using docx2txt (fallback method)"""
        try:
            text_content = docx2txt.process(file_path)
            metadata = {
                'method_used': 'docx2txt',
                'success': False,
                'error': None
            }

            if text_content and text_content.strip():
                metadata['success'] = True
                return {
                    'text': text_content.strip(),
                    'metadata': metadata
                }
            else:
                metadata['error'] = "No text extracted"
                return {'text': "", 'metadata': metadata}

        except Exception as e:
            self.logger.error(f"docx2txt extraction failed: {str(e)}")
            return {
                'text': "",
                'metadata': {
                    'method_used': 'docx2txt',
                    'success': False,
                    'error': str(e)
                }
            }

    def extract_text(self, file_path: str) -> Dict[str, any]:
        """Extract text using primary method, fallback to secondary if needed"""
        if not Path(file_path).exists():
            return {
                'text': "",
                'metadata': {
                    'success': False,
                    'error': "File does not exist"
                }
            }

        # Try python-docx first (primary)
        result = self.extract_text_python_docx(file_path)

        if result['metadata']['success'] and result['text'].strip():
            return result

        # Fallback to docx2txt
        self.logger.info("python-docx failed, trying docx2txt fallback")
        fallback_result = self.extract_text_docx2txt(file_path)

        # Combine metadata from both attempts
        fallback_result['metadata']['primary_method_error'] = result['metadata'].get('error')

        return fallback_result


# Global DOCX parser instance
docx_parser = DOCXParser()